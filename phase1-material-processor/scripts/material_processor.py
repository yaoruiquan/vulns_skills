#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Phase 1 material processor.

This script is intentionally self-contained so the skill can run inside
OpenCode containers and on servers without importing the old vulns project.
"""

from __future__ import annotations

import argparse
import json
import os
import shutil
import sys
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

try:
    import openpyxl
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError as exc:
    print(f"缺少 Python 依赖: {exc}", file=sys.stderr)
    print("请安装 python-docx 和 openpyxl。", file=sys.stderr)
    raise SystemExit(3)


VF_PREFIX_VULN_DESC = "经恒脑AI代码审计智能体分析："
VF_SUFFIX_ANALYSIS = "（查看恒脑AI代码审计智能体官网: https://www.dbappsecurity.com.cn/）"
VF_REPLACE_VERIFICATION = "此分析报告由恒脑AI代码审计智能体自动生成，并经过人工核验。"
DEFAULT_SUBMITTER = "恒脑AI代码审计智能体"


def cell_paragraph_text(cell) -> str:
    return "\n".join(paragraph.text for paragraph in cell.paragraphs)


def _strip_cell_children(cell) -> None:
    """Remove all <w:p> and <w:tbl> children from a table cell."""
    tc = cell._tc
    for item in list(tc.iterchildren()):
        if item.tag.endswith(("p", "tbl")):
            tc.remove(item)


def _cell_has_tables(cell) -> bool:
    """Check if a table cell contains embedded tables."""
    tc = cell._tc
    return any(child.tag.endswith("tbl") for child in tc.iterchildren())


def set_cell_text_plain(cell, text: str) -> None:
    _strip_cell_children(cell)
    new_p = OxmlElement("w:p")
    cell._tc.append(new_p)
    new_r = OxmlElement("w:r")
    new_p.append(new_r)
    new_t = OxmlElement("w:t")
    new_t.text = text
    new_t.set(qn("xml:space"), "preserve")
    new_r.append(new_t)


def insert_prefix_inline(cell, prefix: str) -> None:
    if not cell.paragraphs:
        cell.add_paragraph("")
    paragraph = cell.paragraphs[0]._p
    new_r = OxmlElement("w:r")
    new_t = OxmlElement("w:t")
    new_t.text = prefix
    new_t.set(qn("xml:space"), "preserve")
    new_r.append(new_t)
    paragraph.insert(0, new_r)


def append_suffix_to_cell(cell, suffix: str) -> None:
    paragraph = cell.add_paragraph()
    run = paragraph.add_run()
    run.text = suffix


def ensure_prefix_paragraph(cell, prefix: str) -> bool:
    first_text = cell.paragraphs[0].text if cell.paragraphs else ""
    if first_text.strip().startswith(prefix.strip()):
        return False
    tc = cell._tc
    prefix_p = OxmlElement("w:p")
    tc.insert(0, prefix_p)
    prefix_r = OxmlElement("w:r")
    prefix_p.append(prefix_r)
    prefix_t = OxmlElement("w:t")
    prefix_t.text = prefix
    prefix_t.set(qn("xml:space"), "preserve")
    prefix_r.append(prefix_t)
    return True


def document_first_table(doc_path: Path):
    try:
        doc = Document(str(doc_path))
    except Exception as exc:  # noqa: BLE001
        return None, None, f"无法打开文档: {exc}"
    if not doc.tables:
        return doc, None, "文档中没有表格"
    return doc, doc.tables[0], ""


def modify_cnvd_report(doc_path: Path) -> Tuple[bool, str]:
    doc, table, error = document_first_table(doc_path)
    if error:
        return False, error

    modified_fields: List[str] = []
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        field = cells[0].text.strip()
        target = cells[1]
        text = cell_paragraph_text(target)

        if field == "提交人员":
            if target.text.strip() != DEFAULT_SUBMITTER:
                target.text = DEFAULT_SUBMITTER
            modified_fields.append("提交人员")
        elif field in ("漏洞描述", "漏洞简介"):
            if VF_PREFIX_VULN_DESC in text:
                modified_fields.append(f"{field}(已存在)")
            else:
                insert_prefix_inline(target, VF_PREFIX_VULN_DESC)
                modified_fields.append(field)
        elif field == "漏洞分析":
            changes = []
            if not text.strip().startswith(VF_REPLACE_VERIFICATION):
                ensure_prefix_paragraph(target, VF_REPLACE_VERIFICATION)
                changes.append("开头语句")
            if VF_SUFFIX_ANALYSIS not in text:
                append_suffix_to_cell(target, VF_SUFFIX_ANALYSIS)
                changes.append("后缀")
            modified_fields.append(f"漏洞分析({'+'.join(changes)})" if changes else "漏洞分析(已存在)")
        elif field == "漏洞验证过程":
            if text.strip() == VF_REPLACE_VERIFICATION and not _cell_has_tables(target):
                modified_fields.append("漏洞验证过程(已存在)")
            else:
                set_cell_text_plain(target, VF_REPLACE_VERIFICATION)
                modified_fields.append("漏洞验证过程")

    if modified_fields:
        doc.save(str(doc_path))
        return True, f"已修改字段: {', '.join(modified_fields)}"
    return False, "未找到目标字段"


def modify_cnnvd_report(doc_path: Path) -> Tuple[bool, str]:
    doc, table, error = document_first_table(doc_path)
    if error:
        return False, error

    modified_fields: List[str] = []
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        field = cells[0].text.strip()
        target = cells[1]
        text = cell_paragraph_text(target)

        if field == "提交人员":
            if target.text.strip() != DEFAULT_SUBMITTER:
                target.text = DEFAULT_SUBMITTER
            modified_fields.append("提交人员")
        elif field == "漏洞简介":
            if VF_PREFIX_VULN_DESC in text:
                modified_fields.append("漏洞简介(已存在)")
            else:
                insert_prefix_inline(target, VF_PREFIX_VULN_DESC)
                modified_fields.append("漏洞简介")
        elif field == "漏洞分析":
            if VF_SUFFIX_ANALYSIS in text:
                modified_fields.append("漏洞分析(已存在)")
            else:
                append_suffix_to_cell(target, VF_SUFFIX_ANALYSIS)
                modified_fields.append("漏洞分析")
        elif field == "漏洞验证过程":
            changes = []
            if not text.strip().startswith(VF_REPLACE_VERIFICATION):
                ensure_prefix_paragraph(target, VF_REPLACE_VERIFICATION)
                changes.append("开头语句")
            if VF_SUFFIX_ANALYSIS not in text:
                append_suffix_to_cell(target, VF_SUFFIX_ANALYSIS)
                changes.append("后缀")
            modified_fields.append(f"漏洞验证过程({'+'.join(changes)})" if changes else "漏洞验证过程(已存在)")

    if modified_fields:
        doc.save(str(doc_path))
        return True, f"已修改字段: {', '.join(modified_fields)}"
    return False, "未找到目标字段"


def find_excel_file(data_dir: Path) -> Optional[Path]:
    candidates = [p for p in data_dir.iterdir() if p.is_file()]
    for file_path in candidates:
        if file_path.suffix == ".xlsx" and not file_path.name.startswith((".~", "_", ".")):
            return file_path
    for file_path in data_dir.rglob("*.xlsx"):
        if not file_path.name.startswith((".~", "_", ".")):
            return file_path
    return None


def extract_submitters_from_excel(excel_path: Path) -> Dict[str, str]:
    workbook = openpyxl.load_workbook(excel_path, data_only=True)
    sheet = workbook.active
    header_row_idx = None
    headers: List[str] = []
    for index, row in enumerate(sheet.iter_rows(min_row=1, max_row=10, values_only=True), 1):
        row_text = " ".join(str(cell) for cell in row if cell)
        if "漏洞编号" in row_text or "漏洞标题" in row_text:
            header_row_idx = index
            headers = [str(cell).strip() if cell else "" for cell in row]
            break
    if header_row_idx is None:
        return {}

    das_col = None
    submitter_col = None
    for index, header in enumerate(headers):
        if header == "漏洞编号":
            das_col = index
        if header == "提交人":
            submitter_col = index
    if das_col is None or submitter_col is None:
        return {}

    result: Dict[str, str] = {}
    for row in sheet.iter_rows(min_row=header_row_idx + 1, values_only=True):
        if row[das_col]:
            das_id = str(row[das_col]).strip()
            submitter = str(row[submitter_col]).strip() if row[submitter_col] else ""
            result[das_id] = submitter
    return result


class MaterialService:
    def __init__(self, data_dir: Path):
        self.data_dir = data_dir
        self.template_dirs = [data_dir, data_dir / "test", data_dir / "templates"]
        self._submitter_cache: Optional[Dict[str, str]] = None

    def get_submitters(self) -> Dict[str, str]:
        if self._submitter_cache is None:
            excel_path = find_excel_file(self.data_dir)
            self._submitter_cache = extract_submitters_from_excel(excel_path) if excel_path else {}
        return self._submitter_cache

    def iter_vuln_dirs(self) -> Iterable[Path]:
        seen = set()
        for templates_dir in self.template_dirs:
            if not templates_dir.exists():
                continue
            for item in sorted(templates_dir.iterdir()):
                if item.is_dir() and item.name.startswith("DAS-") and item.name not in seen:
                    seen.add(item.name)
                    yield item

    def find_report_path(self, das_id: str, prefix: str) -> Optional[Path]:
        for vuln_dir in self.iter_vuln_dirs():
            if not vuln_dir.name.startswith(das_id):
                continue
            for subdir in vuln_dir.iterdir():
                if subdir.is_dir() and subdir.name.startswith(prefix):
                    for file_path in sorted(subdir.iterdir()):
                        if file_path.suffix == ".docx" and not file_path.name.startswith("."):
                            return file_path
        return None

    def is_processed(self, doc_path: Optional[Path], report_type: str) -> Optional[bool]:
        if doc_path is None:
            return None
        doc, table, error = document_first_table(doc_path)
        if error:
            return False
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            field = cells[0].text.strip()
            target_text = cell_paragraph_text(cells[1])
            if report_type == "CNVD" and field in ("漏洞描述", "漏洞简介"):
                return VF_PREFIX_VULN_DESC in target_text
            if report_type == "CNNVD" and field == "漏洞简介":
                return VF_PREFIX_VULN_DESC in target_text
        return False

    def process_single(self, das_id: str, submitter: Optional[str] = None) -> Dict[str, Any]:
        if not submitter:
            submitter = self.get_submitters().get(das_id)

        cnvd_path = self.find_report_path(das_id, "CNVD-")
        cnnvd_path = self.find_report_path(das_id, "CNNVD-")
        result: Dict[str, Any] = {
            "das_id": das_id,
            "submitter": submitter or "",
            "success": False,
            "cnvd": {"modified": False, "path": str(cnvd_path) if cnvd_path else None, "message": ""},
            "cnnvd": {"modified": False, "path": str(cnnvd_path) if cnnvd_path else None, "message": ""},
        }

        if cnvd_path:
            modified, message = modify_cnvd_report(cnvd_path)
            result["cnvd"].update({"modified": modified, "message": message})
        else:
            result["cnvd"]["message"] = "未找到 CNVD 报告"

        if cnnvd_path:
            modified, message = modify_cnnvd_report(cnnvd_path)
            result["cnnvd"].update({"modified": modified, "message": message})
        else:
            result["cnnvd"]["message"] = "未找到 CNNVD 报告"

        result["success"] = bool(cnvd_path or cnnvd_path)
        return result

    def process_all(self) -> List[Dict[str, Any]]:
        results = []
        for vuln_dir in self.iter_vuln_dirs():
            das_id = "-".join(vuln_dir.name.split("-")[:2])
            results.append(self.process_single(das_id))
        return results

    def list_vulns(self) -> List[Dict[str, Any]]:
        submitters = self.get_submitters()
        rows = []
        for vuln_dir in self.iter_vuln_dirs():
            das_id = "-".join(vuln_dir.name.split("-")[:2])
            cnvd_path = self.find_report_path(das_id, "CNVD-")
            cnnvd_path = self.find_report_path(das_id, "CNNVD-")
            rows.append({
                "das_id": das_id,
                "folder_name": vuln_dir.name,
                "has_cnvd": cnvd_path is not None,
                "has_cnnvd": cnnvd_path is not None,
                "cnvd_processed": self.is_processed(cnvd_path, "CNVD"),
                "cnnvd_processed": self.is_processed(cnnvd_path, "CNNVD"),
                "submitter": submitters.get(das_id, ""),
            })
        return rows


def count_vuln_dirs(data_dir: Path) -> int:
    return sum(1 for item in data_dir.iterdir() if item.is_dir() and item.name.startswith("DAS-"))


def prepare_work_dir(data_dir: Path, output_root: Optional[Path]) -> Path:
    if output_root is None:
        return data_dir
    vuln_count = count_vuln_dirs(data_dir)
    processed_root = output_root / "processed-materials"
    renamed = processed_root / f"杭州安恒信息原创漏洞报送{vuln_count}个-{data_dir.name}"
    if renamed.exists():
        shutil.rmtree(renamed)
    processed_root.mkdir(parents=True, exist_ok=True)
    shutil.copytree(data_dir, renamed, ignore=shutil.ignore_patterns(".DS_Store", "__MACOSX"))
    return renamed


def render_summary(action: str, work_dir: Path, results: Any) -> str:
    lines = [
        "# phase1-material-processor summary",
        "",
        f"- action: {action}",
        f"- work_dir: {work_dir}",
        "",
    ]
    if action == "list":
        lines.append(f"- vulns: {len(results)}")
        lines.append("")
        for row in results:
            lines.append(
                f"- {row['das_id']}: CNVD={row['has_cnvd']}/{row['cnvd_processed']} "
                f"CNNVD={row['has_cnnvd']}/{row['cnnvd_processed']}"
            )
        return "\n".join(lines) + "\n"

    result_list = results if isinstance(results, list) else [results]
    success_count = sum(1 for item in result_list if item.get("success"))
    cnvd_modified = sum(1 for item in result_list if item.get("cnvd", {}).get("modified"))
    cnnvd_modified = sum(1 for item in result_list if item.get("cnnvd", {}).get("modified"))
    lines.extend([
        f"- processed: {len(result_list)}",
        f"- success: {success_count}/{len(result_list)}",
        f"- cnvd_modified: {cnvd_modified}",
        f"- cnnvd_modified: {cnnvd_modified}",
        "",
    ])
    for item in result_list:
        lines.append(f"## {item.get('das_id')}")
        for key in ("cnvd", "cnnvd"):
            detail = item.get(key, {})
            lines.append(f"- {key.upper()}: modified={detail.get('modified')} path={detail.get('path')}")
            lines.append(f"  message: {detail.get('message')}")
        lines.append("")
    return "\n".join(lines)


def print_list(rows: List[Dict[str, Any]]) -> None:
    print(f"\n发现 {len(rows)} 个漏洞:\n")
    print(f"{'DAS-ID':<15} {'CNVD':<10} {'CNNVD':<10} {'提交人':<10}")
    print("-" * 60)
    for row in rows:
        cnvd = "已处理" if row["cnvd_processed"] else ("未处理" if row["has_cnvd"] else "-")
        cnnvd = "已处理" if row["cnnvd_processed"] else ("未处理" if row["has_cnnvd"] else "-")
        print(f"{row['das_id']:<15} {cnvd:<10} {cnnvd:<10} {row['submitter']:<10}")


def main() -> int:
    parser = argparse.ArgumentParser(description="监管上报前材料整理")
    parser.add_argument("--dir", "-d", required=True, help="输入批次目录，目录下应包含 DAS-* 漏洞目录")
    parser.add_argument("action", nargs="?", default="list", help="操作: batch, list, 或 DAS-ID")
    parser.add_argument("submitter", nargs="?", help="提交人员，可选；默认使用固定值")
    parser.add_argument("--output-root", help="服务化输出根目录；会生成 processed-materials/重命名批次")
    parser.add_argument("--summary", help="summary.txt 输出路径")
    parser.add_argument("--json", dest="json_path", help="JSON 结果输出路径")
    args = parser.parse_args()

    data_dir = Path(args.dir).resolve()
    if not data_dir.is_dir():
        print(f"输入目录不存在: {data_dir}", file=sys.stderr)
        return 2

    output_root = Path(args.output_root).resolve() if args.output_root else None
    work_dir = prepare_work_dir(data_dir, output_root)
    service = MaterialService(work_dir)

    if args.action == "batch":
        results = service.process_all()
        print(f"\n处理了 {len(results)} 个漏洞:\n")
        for item in results:
            cnvd = "✓" if item["cnvd"]["modified"] else "○"
            cnnvd = "✓" if item["cnnvd"]["modified"] else "○"
            print(f"  {item['das_id']}: CNVD={cnvd} CNNVD={cnnvd}")
    elif args.action == "list":
        results = service.list_vulns()
        print_list(results)
    else:
        results = service.process_single(args.action, args.submitter)
        print(json.dumps(results, ensure_ascii=False, indent=2))

    summary = render_summary(args.action, work_dir, results)
    if args.summary:
        summary_path = Path(args.summary).resolve()
        summary_path.parent.mkdir(parents=True, exist_ok=True)
        summary_path.write_text(summary, encoding="utf-8")
    if args.json_path:
        json_path = Path(args.json_path).resolve()
        json_path.parent.mkdir(parents=True, exist_ok=True)
        json_path.write_text(json.dumps(results, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")

    if args.action != "list":
        result_list = results if isinstance(results, list) else [results]
        if not result_list or not any(item.get("success") for item in result_list):
            return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
