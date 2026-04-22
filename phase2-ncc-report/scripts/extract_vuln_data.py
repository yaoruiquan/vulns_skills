#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""从本地漏洞材料中提取 NCC 平台上报所需字段。"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Dict, Iterable, Optional

from docx import Document


SKILL_DIR = Path(__file__).resolve().parents[1]
ENV_FILE = SKILL_DIR / ".env"
DEFAULT_DATA_DIR = os.path.expanduser("~/vulns/date")
MATERIAL_PREFIXES = ("CNVD-", "CNNVD-", "NCC-")
DAS_ID_RE = re.compile(r"(DAS-T\d+)", re.IGNORECASE)
RISK_LEVEL_RE = re.compile(r"风险等级[:：]\s*([^\s]+)")


def load_env() -> None:
    """加载 skill 根目录下的 .env。"""
    if not ENV_FILE.exists():
        return
    for raw_line in ENV_FILE.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        key = key.strip()
        value = value.strip().strip('"').strip("'")
        if key and key not in os.environ:
            os.environ[key] = value


load_env()

PYTHON_PROJECT_PATH = os.environ.get("PYTHON_PROJECT_PATH", "")
if PYTHON_PROJECT_PATH and os.path.isdir(PYTHON_PROJECT_PATH):
    sys.path.insert(0, PYTHON_PROJECT_PATH)

DEFAULT_DATA_DIR = os.environ.get("VULN_DATA_DIR", DEFAULT_DATA_DIR)


def normalize_text(value: str) -> str:
    """清理首尾空白，保留正文换行。"""
    return (value or "").strip()


def first_value(fields: Dict[str, str], *keys: str, default: str = "") -> str:
    """按候选字段名取第一个非空值。"""
    for key in keys:
        value = normalize_text(fields.get(key, ""))
        if value:
            return value
    return default


def detect_material_source(path: Path) -> str:
    """根据目录前缀识别材料来源。"""
    upper_name = path.name.upper()
    for prefix in MATERIAL_PREFIXES:
        if upper_name.startswith(prefix):
            return prefix[:-1]
    return "UNKNOWN"


def extract_das_id(path: Path) -> str:
    """从路径各级目录中提取 DAS-ID。"""
    for part in [path.name, *[parent.name for parent in path.parents]]:
        match = DAS_ID_RE.search(part)
        if match:
            return match.group(1).upper()
    return ""


def is_valid_docx(path: Path) -> bool:
    """过滤掉临时 docx 文件。"""
    if path.suffix.lower() != ".docx":
        return False
    if path.name.startswith("."):
        return False
    if path.name.startswith("~$") or path.name.startswith(".~"):
        return False
    return path.is_file()


def iter_docx_files(folder: Path) -> Iterable[Path]:
    """递归遍历目录中的有效 docx 文件。"""
    for docx_path in sorted(folder.rglob("*.docx")):
        if is_valid_docx(docx_path):
            yield docx_path


def preference_rank(path: Path, prefer_source: str) -> tuple[int, str]:
    """为材料目录或文件计算优先级。"""
    upper_name = path.name.upper()
    prefer_prefix = f"{prefer_source.upper()}-"
    if upper_name.startswith(prefer_prefix):
        return (0, upper_name)
    for index, prefix in enumerate(MATERIAL_PREFIXES, start=1):
        if upper_name.startswith(prefix):
            return (index, upper_name)
    return (len(MATERIAL_PREFIXES) + 1, upper_name)


def find_matching_das_root(das_id: str, data_dir: str) -> Optional[Path]:
    """在漏洞数据根目录下查找匹配 DAS-ID 的目录。"""
    root = Path(data_dir).expanduser()
    if not root.exists():
        return None

    das_id = das_id.upper()
    for item in sorted(root.iterdir()):
        if item.is_dir() and item.name.upper().startswith(das_id):
            return item
    return None


def find_material_dir(root: Path, prefer_source: str) -> Optional[Path]:
    """在输入目录下挑选最合适的材料目录。"""
    if not root.exists():
        return None

    if root.is_file():
        return root.parent if is_valid_docx(root) else None

    if any(is_valid_docx(path) for path in root.glob("*.docx")):
        return root

    candidates = []
    for child in sorted(root.iterdir()):
        if not child.is_dir():
            continue
        if any(is_valid_docx(path) for path in child.glob("*.docx")):
            candidates.append(child)

    if not candidates:
        for docx_path in iter_docx_files(root):
            return docx_path.parent
        return None

    candidates.sort(key=lambda item: preference_rank(item, prefer_source))
    return candidates[0]


def find_preferred_docx(folder: Path) -> Optional[Path]:
    """优先取材料目录根下 docx；没有则递归取第一个。"""
    direct = [path for path in sorted(folder.glob("*.docx")) if is_valid_docx(path)]
    if direct:
        return direct[0]

    for docx_path in iter_docx_files(folder):
        return docx_path
    return None


def list_files(folder: Path, pattern: str) -> list[str]:
    """按 glob 规则列出文件路径。"""
    if not folder.exists():
        return []
    return [str(path) for path in sorted(folder.glob(pattern)) if path.is_file()]


def collect_attachments(material_dir: Path) -> Dict[str, object]:
    """收集材料目录中的附件。"""
    docx_path = find_preferred_docx(material_dir)

    zip_candidates = list_files(material_dir / "poc", "*.zip")
    if not zip_candidates:
        zip_candidates = list_files(material_dir, "*.zip")

    screenshot_paths = []
    screenshot_paths.extend(list_files(material_dir / "poc验证图片", "*"))
    screenshot_paths.extend(list_files(material_dir / "screenshots", "*"))

    video_paths = []
    video_paths.extend(list_files(material_dir / "poc验证视频", "*"))
    video_paths.extend(list_files(material_dir / "videos", "*"))

    upload_files = []
    if docx_path:
        upload_files.append(str(docx_path))
    if zip_candidates:
        upload_files.append(zip_candidates[0])
    upload_files.extend(screenshot_paths)
    upload_files.extend(video_paths)

    return {
        "docx_path": str(docx_path) if docx_path else "",
        "upload_zip_path": zip_candidates[0] if zip_candidates else "",
        "zip_paths": zip_candidates,
        "screenshot_paths": screenshot_paths,
        "video_paths": video_paths,
        "all_upload_files": upload_files,
    }


def extract_fields_from_docx(doc_path: Path) -> Dict[str, str]:
    """从 docx 第一张表提取键值字段。"""
    doc = Document(str(doc_path))
    table = doc.tables[0] if doc.tables else None
    if not table:
        return {}

    fields: Dict[str, str] = {}
    for row in table.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        key = normalize_text(cells[0].text)
        value = normalize_text(cells[1].text)
        if key and value:
            fields[key] = value
    return fields


def extract_risk_level(fields: Dict[str, str]) -> str:
    """尝试从材料文本中提取风险等级。"""
    for key in ("漏洞危害", "漏洞分析", "漏洞描述"):
        match = RISK_LEVEL_RE.search(fields.get(key, ""))
        if match:
            return match.group(1)
    return ""


def extract_impact(fields: Dict[str, str]) -> str:
    """优先返回可直接填表的危害摘要。"""
    direct = first_value(fields, "漏洞危害")
    if direct:
        return direct

    analysis = first_value(fields, "漏洞分析")
    if not analysis:
        return ""

    summary = analysis
    if "漏洞描述" in summary:
        summary = summary.split("漏洞描述", 1)[0]
    summary = summary.replace("此分析报告由恒脑AI代码审计智能体自动生成，并经过人工核验。", "").strip()
    return summary


def resolve_input(args: argparse.Namespace) -> tuple[Optional[Path], Optional[Path], str]:
    """解析脚本输入，得到材料目录和 docx 文件。"""
    prefer_source = args.prefer_source.upper()

    if args.docx_path:
        docx_path = Path(args.docx_path).expanduser()
        if not is_valid_docx(docx_path):
            return None, None, f"无效的 docx 文件: {docx_path}"
        return docx_path.parent, docx_path, ""

    if args.input_path:
        input_path = Path(args.input_path).expanduser()
        if not input_path.exists():
            return None, None, f"输入路径不存在: {input_path}"
        material_dir = find_material_dir(input_path, prefer_source)
        if not material_dir:
            return None, None, f"未在输入路径中找到材料目录: {input_path}"
        docx_path = find_preferred_docx(material_dir)
        if not docx_path:
            return None, None, f"材料目录中未找到 docx: {material_dir}"
        return material_dir, docx_path, ""

    if args.das_id:
        das_root = find_matching_das_root(args.das_id, args.data_dir)
        if not das_root:
            return None, None, f"未找到 DAS 目录: {args.das_id}"
        material_dir = find_material_dir(das_root, prefer_source)
        if not material_dir:
            return None, None, f"未在 DAS 目录中找到材料目录: {das_root}"
        docx_path = find_preferred_docx(material_dir)
        if not docx_path:
            return None, None, f"材料目录中未找到 docx: {material_dir}"
        return material_dir, docx_path, ""

    return None, None, "缺少输入参数，请提供 DAS-ID、--input-path 或 --docx-path"


def extract_ncc_data(material_dir: Path, docx_path: Path) -> Dict[str, object]:
    """提取 NCC 平台上报所需的统一数据。"""
    fields = extract_fields_from_docx(docx_path)
    attachments = collect_attachments(material_dir)
    das_root = material_dir.parent if material_dir.parent != material_dir else material_dir
    das_id = extract_das_id(material_dir) or extract_das_id(docx_path)

    data: Dict[str, object] = {
        "platform": "NCC",
        "das_id": das_id,
        "input_root": str(das_root),
        "material_dir": str(material_dir),
        "material_source": detect_material_source(material_dir),
        "docx_path": attachments["docx_path"],
        "upload_zip_path": attachments["upload_zip_path"],
        "zip_paths": attachments["zip_paths"],
        "screenshot_paths": attachments["screenshot_paths"],
        "video_paths": attachments["video_paths"],
        "all_upload_files": attachments["all_upload_files"],
        "title": first_value(fields, "漏洞名称"),
        "description": first_value(fields, "漏洞描述", "漏洞简介"),
        "impact": extract_impact(fields),
        "vuln_type": first_value(fields, "漏洞类型"),
        "target_type": first_value(fields, "影响对象类型"),
        "url": first_value(fields, "漏洞URL", "漏洞定位"),
        "unit_name": first_value(fields, "漏洞厂商", "影响厂商", "厂商名称", "提交机构"),
        "vendor_website": first_value(fields, "厂商官网"),
        "affected_product": first_value(fields, "影响产品", "受影响实体", "产品名称"),
        "version": first_value(fields, "影响版本", "受影响实体版本号", "版本号"),
        "request_method": first_value(fields, "请求方式"),
        "discoverer_name": first_value(fields, "提交人员", "发现者"),
        "contact": first_value(fields, "联系方式", "联系电话"),
        "submit_org": first_value(fields, "提交机构"),
        "submit_date": first_value(fields, "提交日期"),
        "verification": first_value(fields, "漏洞验证过程", "验证过程", "漏洞验证"),
        "temporary_solution": first_value(fields, "临时解决方案", default="无"),
        "formal_solution": first_value(fields, "正式解决方案", "修复方案", default="见附件"),
        "risk_level": extract_risk_level(fields),
    }

    return data


def build_parser() -> argparse.ArgumentParser:
    """构造命令行参数解析器。"""
    parser = argparse.ArgumentParser(description="提取 NCC 平台上报字段")
    parser.add_argument("das_id", nargs="?", help="DAS-ID，兼容旧用法")
    parser.add_argument("--platform", default="NCC", help="保留兼容参数，当前仅支持 NCC")
    parser.add_argument("--data-dir", default=DEFAULT_DATA_DIR, help="漏洞数据根目录")
    parser.add_argument("--input-path", default="", help="具体 DAS 目录或材料目录")
    parser.add_argument("--docx-path", default="", help="直接指定 docx 文件路径")
    parser.add_argument(
        "--prefer-source",
        default="CNVD",
        choices=["CNVD", "CNNVD", "NCC"],
        help="同一 DAS 目录下优先选择哪类材料目录，默认 CNVD",
    )
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()

    material_dir, docx_path, error = resolve_input(args)
    if error:
        print(json.dumps({"error": error}, ensure_ascii=False, indent=2))
        return 1

    assert material_dir is not None
    assert docx_path is not None
    data = extract_ncc_data(material_dir, docx_path)
    print(json.dumps(data, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
