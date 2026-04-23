#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""从 docx 提取漏洞数据，供 CNVD/CNNVD 上报使用"""

import sys
import os
import json
import re
from pathlib import Path

# 加载 .env 配置
SKILL_DIR = Path(__file__).parent.parent
ENV_FILE = SKILL_DIR / ".env"

if ENV_FILE.exists():
    with open(ENV_FILE) as f:
        for line in f:
            line = line.strip()
            if line and not line.startswith("#") and "=" in line:
                key, val = line.split("=", 1)
                os.environ[key.strip()] = val.strip()

# Python 项目路径（可选）
PYTHON_PROJECT_PATH = os.environ.get("PYTHON_PROJECT_PATH", "")
if PYTHON_PROJECT_PATH and os.path.isdir(PYTHON_PROJECT_PATH):
    sys.path.insert(0, PYTHON_PROJECT_PATH)

from docx import Document
from typing import Dict, Optional

# 数据目录（从环境变量读取，或使用默认值）
DEFAULT_DATA_DIR = os.environ.get("VULN_DATA_DIR", os.path.expanduser("~/vulns/date"))


def find_docx_path(das_id: str, platform: str, data_dir: str = DEFAULT_DATA_DIR) -> Optional[str]:
    """查找漏洞 docx 文件路径"""
    for item in os.listdir(data_dir):
        item_path = os.path.join(data_dir, item)
        if item.startswith(das_id) and os.path.isdir(item_path):
            vuln_folder = item_path
            # 查找平台子目录
            for sub in os.listdir(vuln_folder):
                sub_path = os.path.join(vuln_folder, sub)
                if sub.startswith(f"{platform}-") and os.path.isdir(sub_path):
                    platform_folder = sub_path
                    for f in os.listdir(platform_folder):
                        f_path = os.path.join(platform_folder, f)
                        if f.endswith('.docx') and not f.startswith('.') and os.path.isfile(f_path):
                            return f_path
    return None


def find_attachment_zip_path(platform_folder: str, platform: str) -> str:
    """查找平台原始 zip 附件，优先使用 CNVD/CNNVD 整包 zip，不重新压缩目录"""
    folder = Path(platform_folder)
    if not folder.exists():
        return ""

    platform_upper = platform.upper()
    search_roots = [folder.parent, folder]

    platform_zips = []
    fallback_zips = []
    for root in search_roots:
        if not root.exists():
            continue
        for path in root.iterdir():
            if not path.is_file() or path.suffix.lower() != ".zip" or path.name.startswith("."):
                continue
            if path.name.upper().startswith(platform_upper):
                platform_zips.append(path)
            elif root == folder:
                fallback_zips.append(path)

    zip_files = platform_zips or fallback_zips
    if not zip_files:
        return ""

    return str(max(zip_files, key=lambda path: path.stat().st_size))


def clean_cnvd_description(description: str) -> str:
    """清理 CNVD 漏洞描述中不应填写到表单的固定分析前缀"""
    if not description:
        return ""

    cleaned = description.strip()
    cleaned = re.sub(
        r"^\s*经恒脑\s*AI\s*代码审计智能体分析[:：]\s*",
        "",
        cleaned,
        count=1,
    )
    return cleaned.strip()


def extract_fields_from_docx(doc_path: str) -> Dict[str, str]:
    """从 docx 提取所有字段"""
    doc = Document(doc_path)
    table = doc.tables[0] if doc.tables else None
    if not table:
        return {}

    fields = {}
    for row in table.rows:
        cells = row.cells
        if len(cells) >= 2:
            key = cells[0].text.strip()
            val = cells[1].text.strip()
            if key and val:
                fields[key] = val
    return fields


def map_cnvd_vuln_type(vuln_type_text: str) -> str:
    """将漏洞类型文本映射为 CNVD 表单值"""
    text = (vuln_type_text or "").strip()
    if not text:
        return "其他"

    # 先匹配更具体的高优先级类型，避免“内存/缓冲区”等泛化词提前落到“其他”。
    mapping = [
        ("二进制", "二进制"),
        ("内存缓冲区", "二进制"),
        ("缓冲区", "二进制"),
        ("堆溢出", "二进制"),
        ("栈溢出", "二进制"),
        ("溢出", "二进制"),
        ("越界", "二进制"),
        ("释放后使用", "二进制"),
        ("UAF", "二进制"),
        ("uaf", "二进制"),
        ("SQL注入", "sql注入"),
        ("sql注入", "sql注入"),
        ("XSS", "跨站脚本"),
        ("跨站脚本", "跨站脚本"),
        ("文件上传", "文件上传"),
        ("信息泄露", "信息泄露"),
        ("SSRF", "其他"),
        ("弱口令", "其他"),
        ("未授权访问", "其他"),
        ("逻辑缺陷", "其他"),
        ("文件包含", "其他"),
        ("命令执行", "其他"),
        ("代码执行", "其他"),
        ("反序列化", "其他"),
        ("目录遍历", "其他"),
    ]
    for key, val in mapping:
        if key in text:
            return val
    return "其他"


def map_soft_style(obj_type_text: str) -> str:
    """将影响对象类型映射为 CNVD 表单值"""
    mapping = {
        "WEB应用": "29",
        "操作系统": "27",
        "应用程序": "28",
        "数据库": "30",
        "网络设备": "31",
        "安全产品": "32",
        "智能设备": "33",
        "工业控制": "38",
    }
    for key, val in mapping.items():
        if key in obj_type_text:
            return val
    return "28"


def extract_cnvd_data(das_id: str, data_dir: str = DEFAULT_DATA_DIR) -> Dict[str, str]:
    """提取 CNVD 上报所需数据"""
    doc_path = find_docx_path(das_id, "CNVD", data_dir)
    if not doc_path:
        return {"error": f"未找到 CNVD docx: {das_id}"}

    fields = extract_fields_from_docx(doc_path)
    folder_path = os.path.dirname(doc_path)
    description = clean_cnvd_description(fields.get("漏洞描述", ""))
    attachment_zip_path = find_attachment_zip_path(folder_path, "CNVD")

    return {
        "das_id": das_id,
        "title": fields.get("漏洞名称", ""),
        "description": description,
        "vuln_type": map_cnvd_vuln_type(fields.get("漏洞类型", "")),
        "vuln_type_raw": fields.get("漏洞类型", ""),
        "url": fields.get("漏洞URL", ""),
        "unit_name": fields.get("漏洞厂商", ""),
        "is_event": "0",  # 默认通用型
        "soft_style_id": map_soft_style(fields.get("影响对象类型", "")),
        "discoverer_name": fields.get("提交人员", ""),
        "affected_product": fields.get("影响产品", ""),
        "version": fields.get("影响版本", ""),
        "folder_path": folder_path,
        "docx_path": doc_path,
        "attachment_zip_path": attachment_zip_path,
    }


def extract_cnnvd_data(das_id: str, data_dir: str = DEFAULT_DATA_DIR) -> Dict[str, str]:
    """提取 CNNVD 上报所需数据"""
    doc_path = find_docx_path(das_id, "CNNVD", data_dir)
    if not doc_path:
        return {"error": f"未找到 CNNVD docx: {das_id}"}

    fields = extract_fields_from_docx(doc_path)
    folder_path = os.path.dirname(doc_path)
    attachment_zip_path = find_attachment_zip_path(folder_path, "CNNVD")

    return {
        "das_id": das_id,
        "title": fields.get("漏洞名称", ""),
        "description": fields.get("漏洞简介", ""),
        "vuln_type": fields.get("漏洞类型", ""),
        "url": fields.get("漏洞定位", ""),
        "affected_product": fields.get("受影响实体版本号", ""),
        "discoverer_name": fields.get("提交人员", ""),
        "contact": fields.get("联系方式", ""),
        "verification": fields.get("漏洞验证过程", ""),
        "folder_path": folder_path,
        "docx_path": doc_path,
        "attachment_zip_path": attachment_zip_path,
    }


def main():
    if len(sys.argv) < 2:
        print("Usage: python3 extract_vuln_data.py <das_id> [--platform CNVD|CNNVD] [--data-dir <path>]")
        print("  默认平台: CNVD")
        sys.exit(1)

    das_id = sys.argv[1]
    platform = "CNVD"
    data_dir = DEFAULT_DATA_DIR

    # 解析参数
    i = 2
    while i < len(sys.argv):
        if sys.argv[i] == "--platform" and i + 1 < len(sys.argv):
            platform = sys.argv[i + 1].upper()
            i += 2
        elif sys.argv[i] == "--data-dir" and i + 1 < len(sys.argv):
            data_dir = sys.argv[i + 1]
            i += 2
        else:
            i += 1

    # 提取数据
    if platform == "CNNVD":
        data = extract_cnnvd_data(das_id, data_dir)
    else:
        data = extract_cnvd_data(das_id, data_dir)

    # 输出 JSON
    print(json.dumps(data, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
